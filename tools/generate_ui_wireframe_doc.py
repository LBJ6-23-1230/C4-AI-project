from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\zhixue_Mate")
OUT_DIR = ROOT / "deliverables" / "ui_wireframes"
DOCX_PATH = OUT_DIR / "知学搭子_前端页面草图设计.docx"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")


COLORS = {
    "ink": "#1F2937",
    "muted": "#6B7280",
    "line": "#CBD5E1",
    "panel": "#F8FAFC",
    "accent": "#2563EB",
    "soft_accent": "#DBEAFE",
    "green": "#10B981",
    "soft_green": "#D1FAE5",
    "amber": "#F59E0B",
    "soft_amber": "#FEF3C7",
}


def rgb(hex_color: str):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False):
    return ImageFont.truetype(str(FONT_PATH), size=size, index=0)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, style=None, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    return p


def draw_text(draw, xy, text, fnt, fill="#1F2937", max_width=None, line_gap=6):
    x, y = xy
    if max_width is None:
        draw.text((x, y), text, font=fnt, fill=rgb(fill))
        return y + draw.textbbox((x, y), text, font=fnt)[3] - y

    lines = []
    current = ""
    for ch in text:
        trial = current + ch
        width = draw.textbbox((0, 0), trial, font=fnt)[2]
        if width <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)

    for line in lines:
        draw.text((x, y), line, font=fnt, fill=rgb(fill))
        y += fnt.size + line_gap
    return y


def rounded(draw, box, fill="#FFFFFF", outline="#CBD5E1", width=2, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill), outline=rgb(outline), width=width)


def pill(draw, box, text, fill="#DBEAFE", text_fill="#2563EB"):
    rounded(draw, box, fill=fill, outline=fill, width=1, radius=20)
    x1, y1, x2, y2 = box
    f = font(24)
    tb = draw.textbbox((0, 0), text, font=f)
    draw.text((x1 + (x2 - x1 - (tb[2] - tb[0])) / 2, y1 + (y2 - y1 - (tb[3] - tb[1])) / 2 - 2),
              text, font=f, fill=rgb(text_fill))


def card(draw, x, y, w, h, title, lines, accent="#2563EB"):
    rounded(draw, (x, y, x + w, y + h), fill="#FFFFFF", outline="#D7DEE8", width=2, radius=24)
    draw.rounded_rectangle((x, y, x + 10, y + h), radius=8, fill=rgb(accent))
    draw_text(draw, (x + 28, y + 22), title, font(30, True), COLORS["ink"], max_width=w - 56)
    cy = y + 70
    for line in lines:
        cy = draw_text(draw, (x + 28, cy), line, font(23), COLORS["muted"], max_width=w - 56)
        cy += 4


def create_phone_canvas(title):
    img = Image.new("RGB", (760, 1320), rgb("#EEF2F7"))
    draw = ImageDraw.Draw(img)
    rounded(draw, (38, 28, 722, 1292), fill="#FFFFFF", outline="#AAB5C4", width=4, radius=48)
    draw.rounded_rectangle((305, 48, 455, 62), radius=8, fill=rgb("#CBD5E1"))
    draw_text(draw, (82, 94), title, font(38, True), COLORS["ink"])
    return img, draw


def wire_home():
    img, draw = create_phone_canvas("首页")
    draw_text(draw, (82, 150), "你好，小明  ·  今日 7月10日", font(24), COLORS["muted"])
    card(draw, 82, 205, 596, 210, "AI 今日学习重点", [
        "数据结构 · 优先级高",
        "距考试 5 天｜近期复习不足",
        "今晚 20:00-22:00 复习二叉树遍历",
    ], COLORS["accent"])
    labels = ["今日建议", "拍照错题", "学习标签", "找搭子"]
    for i, label in enumerate(labels):
        x = 82 + (i % 2) * 306
        y = 455 + (i // 2) * 102
        rounded(draw, (x, y, x + 276, y + 76), fill="#F8FAFC", outline="#D7DEE8", radius=20)
        draw_text(draw, (x + 70, y + 23), label, font(25, True), COLORS["ink"])
    card(draw, 82, 700, 596, 185, "学习概览", [
        "目标：数据结构期末 80+",
        "薄弱点：二叉树遍历 / 递归理解",
        "优势点：图算法基础",
    ], COLORS["green"])
    card(draw, 82, 920, 596, 165, "最近错题", [
        "二叉树遍历 · 递归理解不足",
        "已生成标签：二叉树薄弱",
    ], COLORS["amber"])
    nav(draw)
    return img


def wire_suggestion():
    img, draw = create_phone_canvas("今日学习建议")
    card(draw, 82, 165, 596, 225, "AI 推荐任务卡", [
        "数据结构｜优先级：高",
        "原因：距考试 5 天、近期复习不足",
        "错题集中在二叉树遍历",
    ], COLORS["accent"])
    draw_text(draw, (82, 430), "今晚计划", font(30, True), COLORS["ink"])
    steps = [
        ("20:00", "二叉树遍历复习"),
        ("20:30", "递归过程练习"),
        ("21:00", "完成 2 道基础题"),
        ("21:20", "简短复盘"),
    ]
    y = 485
    for tm, text in steps:
        pill(draw, (82, y, 178, y + 42), tm, fill=COLORS["soft_accent"])
        draw_text(draw, (205, y + 7), text, font(24), COLORS["ink"])
        y += 76
    card(draw, 82, 805, 596, 150, "关联信息", [
        "相关错题：2 道",
        "相关标签：二叉树薄弱 / 递归理解待提升",
    ], COLORS["green"])
    bottom_buttons(draw, ["开始学习", "调整时间", "找搭子"])
    return img


def wire_wrong_question():
    img, draw = create_phone_canvas("错题分析")
    rounded(draw, (82, 165, 678, 395), fill="#F8FAFC", outline="#CBD5E1", radius=24)
    draw_text(draw, (275, 250), "错题图片预览", font(28, True), COLORS["muted"])
    card(draw, 82, 430, 596, 215, "诊断卡", [
        "课程：数据结构",
        "知识点：二叉树遍历",
        "错误类型：递归理解不足",
    ], COLORS["accent"])
    card(draw, 82, 685, 596, 190, "原因分析", [
        "对前序、中序、后序遍历的递归过程掌握不牢。",
        "建议复习三种遍历方式并完成 2 道基础练习。",
    ], COLORS["amber"])
    draw_text(draw, (82, 920), "生成标签", font(30, True), COLORS["ink"])
    pill(draw, (82, 975, 260, 1025), "二叉树薄弱", fill=COLORS["soft_accent"])
    pill(draw, (285, 975, 548, 1025), "递归理解待提升", fill=COLORS["soft_amber"], text_fill="#92400E")
    bottom_buttons(draw, ["加入计划", "生成相似题", "找搭子"])
    return img


def wire_tags():
    img, draw = create_phone_canvas("学习标签")
    card(draw, 82, 165, 596, 140, "目标卡", ["数据结构期末 80+"], COLORS["accent"])
    draw_text(draw, (82, 345), "薄弱标签", font(30, True), COLORS["ink"])
    pill(draw, (82, 395, 250, 445), "二叉树遍历", fill=COLORS["soft_amber"], text_fill="#92400E")
    pill(draw, (272, 395, 440, 445), "递归理解", fill=COLORS["soft_amber"], text_fill="#92400E")
    pill(draw, (462, 395, 665, 445), "前中后序遍历", fill=COLORS["soft_amber"], text_fill="#92400E")
    draw_text(draw, (82, 495), "优势标签", font(30, True), COLORS["ink"])
    pill(draw, (82, 545, 260, 595), "图算法基础", fill=COLORS["soft_green"], text_fill="#047857")
    card(draw, 82, 650, 596, 235, "标签来源", [
        "二叉树遍历：来自错题 q001",
        "错误类型：递归理解不足",
        "建议：完成基础练习 2 道",
    ], COLORS["green"])
    card(draw, 82, 925, 596, 150, "学习画像", [
        "考试临近｜数据结构优先｜需要概念补强",
    ], COLORS["accent"])
    bottom_buttons(draw, ["更新目标", "相关错题", "匹配搭子"])
    return img


def wire_match():
    img, draw = create_phone_canvas("学习搭子匹配")
    card(draw, 82, 165, 596, 165, "最佳匹配：小李", [
        "匹配度 92｜同专业大二｜今晚可学",
    ], COLORS["accent"])
    draw_text(draw, (82, 365), "匹配理由", font(30, True), COLORS["ink"])
    reasons = [("课程一致", "+40"), ("目标一致", "+25"), ("知识互补", "+20"), ("时间重合", "+15")]
    y = 415
    for label, score in reasons:
        rounded(draw, (82, y, 678, y + 58), fill="#F8FAFC", outline="#D7DEE8", radius=16)
        draw_text(draw, (112, y + 15), label, font(24), COLORS["ink"])
        draw_text(draw, (600, y + 15), score, font(24, True), COLORS["accent"])
        y += 76
    card(draw, 82, 740, 596, 165, "互补关系", [
        "你薄弱：二叉树遍历｜TA 擅长：二叉树遍历",
        "你擅长：图算法基础｜TA 薄弱：图算法",
    ], COLORS["green"])
    card(draw, 82, 945, 596, 160, "30 分钟协同计划", [
        "20:00-20:10 互讲遍历",
        "20:10-20:20 互相出题测试",
        "20:20-20:30 错题互评与总结",
    ], COLORS["amber"])
    bottom_buttons(draw, ["邀请搭子", "换一个", "学习房间"])
    return img


def nav(draw):
    rounded(draw, (82, 1145, 678, 1225), fill="#F8FAFC", outline="#D7DEE8", radius=24)
    for i, label in enumerate(["首页", "建议", "错题", "标签", "搭子"]):
        draw_text(draw, (118 + i * 112, 1172), label, font(22), COLORS["muted"])


def bottom_buttons(draw, labels):
    y = 1160
    widths = [170, 190, 170]
    x = 82
    for i, label in enumerate(labels):
        w = widths[i]
        fill = COLORS["accent"] if i == 0 else "#F8FAFC"
        text_fill = "#FFFFFF" if i == 0 else COLORS["ink"]
        rounded(draw, (x, y, x + w, y + 64), fill=fill, outline="#D7DEE8", radius=22)
        draw_text(draw, (x + 25, y + 18), label, font(23, True), text_fill)
        x += w + 24


def save_wireframes():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    makers = [
        ("01_home.png", wire_home),
        ("02_suggestion.png", wire_suggestion),
        ("03_wrong_question.png", wire_wrong_question),
        ("04_tags.png", wire_tags),
        ("05_match.png", wire_match),
    ]
    paths = []
    for name, maker in makers:
        path = OUT_DIR / name
        maker().save(path, quality=95)
        paths.append(path)
    return paths


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 18, "2563EB"),
        ("Heading 2", 14, "1F2937"),
        ("Heading 3", 12, "1F2937"),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def build_doc(wireframe_paths):
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("知学搭子前端页面草图设计")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1F2937")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("基于 AI 流程图与数据结构设计的 5 个核心页面方案")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("6B7280")

    doc.add_heading("总体设计思路", level=1)
    add_para(doc, "产品主线：课程 / DDL / 考试输入 → Agent 判断学习优先级 → 今日学习建议 → 错题诊断 → 学习标签沉淀 → 学习搭子匹配 → 30 分钟协同学习计划。")
    add_bullets(doc, [
        "首页回答“今天最该学什么”，用 AI 今日重点卡承接 course.priority、examDaysLeft 和 recentStatus。",
        "今日学习建议页强调推荐原因和可执行时间块，让用户从建议直接进入学习。",
        "错题分析页按“错题事实 → 知识诊断 → 补强动作”组织，输出 wrongQuestion 和 tags。",
        "学习标签页把 weakness、strength 与标签来源可视化，形成用户学习画像。",
        "搭子匹配页突出 matchResult 的证据：课程一致、目标一致、知识互补、时间重合，并给出 30 分钟计划。",
    ])

    doc.add_heading("页面关系", level=1)
    flow_table = doc.add_table(rows=1, cols=5)
    flow_table.autofit = False
    labels = ["首页", "今日建议", "错题分析", "学习标签", "搭子匹配"]
    for cell, label in zip(flow_table.rows[0].cells, labels):
        set_cell_shading(cell, "DBEAFE")
        set_cell_text(cell, label, bold=True, color="1F2937")
    add_para(doc, "交互上建议保留底部导航：首页｜建议｜错题｜标签｜搭子；核心任务流也可以从首页 AI 今日重点卡连续进入。")

    pages = [
        ("1. 首页", "学习状态总览 + 进入核心流程", [
            "首屏放 AI 今日学习重点卡，显示课程、优先级、考试倒计时和推荐时间。",
            "提供今日建议、拍照错题、学习标签、找搭子四个快捷入口。",
            "下方展示目标、薄弱点、优势点与最近错题，形成学习闭环入口。",
        ]),
        ("2. 今日学习建议页", "Agent 输出今日任务，帮助用户马上开始", [
            "建议卡展示推荐课程、优先级和推荐原因。",
            "计划区用时间轴表达 20:00-21:30 的学习安排。",
            "关联错题和标签，让建议看起来不是凭空生成。",
        ]),
        ("3. 错题分析页", "拍照上传错题后输出知识点、错误原因和补强建议", [
            "顶部保留错题图片预览区，增强识别可信度。",
            "诊断卡展示课程、知识点、错误类型。",
            "底部生成标签，并提供加入今日计划、生成相似题、找搭子的动作。",
        ]),
        ("4. 学习标签页", "展示 AI 对用户的学习画像理解", [
            "目标卡展示用户当前学习目标。",
            "薄弱标签和优势标签分区，便于后续互补匹配。",
            "标签来源区说明标签来自错题、计划或学习记录，避免标签显得随意。",
        ]),
        ("5. 学习搭子匹配页", "展示匹配对象、匹配理由和协同学习计划", [
            "最佳匹配卡展示对象和匹配度。",
            "匹配理由拆成四项计分：课程一致 40、目标一致 25、知识互补 20、时间重合 15。",
            "协同学习计划直接给出 30 分钟步骤，降低沟通成本。",
        ]),
    ]

    for idx, (heading, role, bullets) in enumerate(pages):
        doc.add_page_break()
        doc.add_heading(heading, level=1)
        add_para(doc, f"页面定位：{role}", bold=True, color="2563EB")
        doc.add_heading("草图", level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(wireframe_paths[idx]), width=Inches(3.05))
        doc.add_heading("设计说明", level=2)
        add_bullets(doc, bullets)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def main():
    paths = save_wireframes()
    build_doc(paths)
    print(DOCX_PATH)
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
